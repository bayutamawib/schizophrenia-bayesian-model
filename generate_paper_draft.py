"""
Academic Paper Draft Generator
===============================
Generates a formatted academic paper draft in paper_draft.docx
using arXiv formatting conventions and python-docx.

Content synthesized from:
- Ford et al. (2014) - Main study
- Pynn & DeSouza (2012) - Efference copy review
- Adams & Friston (2016) - Bayesian brain in schizophrenia
- Fletcher & Frith (2008) - Bayesian approach to positive symptoms
- Valton et al. (2017) - Computational modelling review
- Our calculated Bayesian model results
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


# ============================================================
# FORMATTING HELPERS
# ============================================================

def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shd)


def add_horizontal_rule(doc, thickness=1):
    """Add a horizontal rule to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): str(thickness * 8),
        qn('w:space'): '1',
        qn('w:color'): '000000'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_paragraph_font(paragraph, font_name='Times New Roman',
                       font_size=10, bold=False, italic=False,
                       color=None, alignment=None):
    """Configure paragraph formatting."""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
    paragraph.paragraph_format.space_after = Pt(5.5)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(11)
    if alignment:
        paragraph.alignment = alignment


def add_body_text(doc, text):
    """Add a body paragraph with standard formatting."""
    p = doc.add_paragraph(text)
    set_paragraph_font(p, font_size=10)
    p.paragraph_format.first_line_indent = Pt(0)
    return p


def add_heading_h1(doc, text):
    """Add H1: ALL CAPS, flush left, bold, 12pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(5.5)
    run = p.add_run(text.upper())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    return p


def add_heading_h2(doc, text):
    """Add H2: Initial caps, flush left, bold, 10pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(5.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.bold = True
    return p


def add_heading_h3(doc, text):
    """Add H3: Flush left, initial caps, bold, 10pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(5.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    return p


def add_table(doc, headers, rows, caption=None):
    """Add a formatted table with caption."""
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(11)
        run = p.add_run(caption)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.font.bold = True

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)
                run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing after table
    return table


# ============================================================
# PAPER CONTENT
# ============================================================

def write_title_block(doc):
    """Write title, authors, and affiliations."""
    add_horizontal_rule(doc, thickness=4)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "A Bayesian Inference Model of Corollary Discharge Dysfunction "
        "in Schizophrenia: Quantifying Efference Copy Precision from EEG Data"
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True

    add_horizontal_rule(doc, thickness=1)

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("[Author Name]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

    # Affiliation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Institution, Department]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True

    # Email
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(11)
    run = p.add_run("[email@institution.edu]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)


def write_abstract(doc):
    """Write the abstract section."""
    add_heading_h1(doc, "Abstract")

    add_body_text(doc,
        "Predictive coding theories propose that the brain operates as a Bayesian "
        "inference engine, continuously generating predictions about sensory input and "
        "updating beliefs based on prediction errors. In the context of self-generated "
        "actions, the motor system sends an efference copy of the motor command to "
        "sensory cortex, producing a corollary discharge that suppresses neural "
        "responses to expected sensory consequences. Dysfunction in this mechanism "
        "has been consistently implicated in the pathophysiology of schizophrenia, "
        "particularly in explaining auditory hallucinations and delusions of agency. "
        "In this study, we formalize the corollary discharge process as a linear "
        "Bayesian model and apply it to EEG data from 81 participants (32 healthy "
        "controls, 49 schizophrenia patients). Our model quantifies the precision "
        "weight (x) with which the brain integrates motor preparation signals into "
        "auditory predictions. Results show that while the efference copy signal "
        "(C3_B1) is approximately 45% weaker in schizophrenia patients (0.79 vs "
        "1.44 µV), the computed precision weight is paradoxically higher (-1.89 vs "
        "-1.40), suggesting a compensatory over-weighting mechanism. Additionally, "
        "an ANCOVA-based variance proxy analysis was conducted to approximate "
        "inter-trial phase coherence (ITPC), revealing directionally correct but "
        "non-significant differences in efference copy consistency between groups. "
        "These findings support the Bayesian brain hypothesis of schizophrenia and "
        "provide a quantitative framework for characterizing predictive coding "
        "deficits from standard EEG measurements."
    )

    # Keywords
    p = doc.add_paragraph()
    run_bold = p.add_run("Keywords: ")
    run_bold.font.name = 'Times New Roman'
    run_bold.font.size = Pt(10)
    run_bold.font.bold = True
    run = p.add_run(
        "Bayesian inference, predictive coding, efference copy, corollary "
        "discharge, schizophrenia, EEG, N100, precision weighting"
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True


def write_introduction(doc):
    """Write the introduction section."""
    add_heading_h1(doc, "1. Introduction")

    add_heading_h2(doc, "1.1 The Bayesian Brain Hypothesis")
    add_body_text(doc,
        "Contemporary neuroscience increasingly views the brain as a "
        "probabilistic inference machine that constructs internal models of the "
        "world and continuously updates them based on incoming sensory evidence "
        "(Adams and Friston, 2016). Under this framework, known as the Bayesian "
        "brain hypothesis, perception is not a passive registration of sensory "
        "input but an active process of hypothesis testing, where prior beliefs "
        "are combined with sensory likelihood to form posterior beliefs about the "
        "causes of sensory data (Fletcher and Frith, 2008)."
    )
    add_body_text(doc,
        "A critical feature of Bayesian inference is the concept of precision — "
        "the inverse variance of a probability distribution — which determines "
        "the relative weighting of prior expectations versus sensory evidence. "
        "When priors are highly precise (confident), the brain relies more on "
        "its predictions; when sensory data are highly precise, predictions are "
        "updated more readily. This precision-weighting mechanism provides a "
        "formal computational account of how the brain balances top-down "
        "expectations with bottom-up sensory signals (Adams and Friston, 2016)."
    )

    add_heading_h2(doc, "1.2 Efference Copy and Corollary Discharge")
    add_body_text(doc,
        "A well-characterized example of predictive coding in action is the "
        "efference copy / corollary discharge system. When the brain initiates a "
        "motor action, a copy of the motor command (the efference copy) is sent "
        "from motor areas to the relevant sensory cortex. This efference copy "
        "generates a corollary discharge — a neural prediction of the expected "
        "sensory consequence of the action (Pynn and DeSouza, 2012). The "
        "predicted sensation is then compared with the actual sensory input, and "
        "if they match, the sensory response is suppressed. This is why we "
        "cannot tickle ourselves: our brain accurately predicts the tactile "
        "sensation and cancels it out."
    )
    add_body_text(doc,
        "This mechanism has been extensively documented across species and "
        "sensory modalities. In the auditory domain, the N1 (or N100) component "
        "of the event-related potential (ERP) — a negative deflection peaking "
        "approximately 100 milliseconds after stimulus onset — is reliably "
        "suppressed when the auditory stimulus is self-generated compared to when "
        "it is externally generated (Ford et al., 2014). The lateralized readiness "
        "potential (LRP), measured over motor cortex prior to a voluntary movement, "
        "provides a neural index of the efference copy itself (Ford et al., 2014; "
        "Pynn and DeSouza, 2012)."
    )

    add_heading_h2(doc, "1.3 Predictive Coding Dysfunction in Schizophrenia")
    add_body_text(doc,
        "Schizophrenia is a severe psychiatric disorder characterized by positive "
        "symptoms including hallucinations and delusions, as well as negative "
        "symptoms such as avolition and alogia (Fletcher and Frith, 2008). A "
        "growing body of evidence suggests that many of these symptoms arise from "
        "fundamental disruptions in predictive coding mechanisms. Specifically, "
        "patients with schizophrenia may be unable to accurately predict the "
        "sensory consequences of their own actions, leading self-generated "
        "sensations to be experienced as externally caused — a plausible "
        "computational basis for auditory verbal hallucinations."
    )
    add_body_text(doc,
        "Ford et al. (2014) demonstrated that patients with schizophrenia show "
        "significantly reduced N1 suppression when pressing a button to deliver "
        "a tone, indicating impaired corollary discharge. Furthermore, the "
        "lateralized readiness potential (LRP) preceding the button press was "
        "smaller in patients, suggesting that the efference copy signal itself "
        "is degraded at its source. Critically, LRP amplitude was correlated "
        "with N1 suppression in both groups, confirming that stronger motor "
        "planning signals produce stronger sensory predictions."
    )
    add_body_text(doc,
        "From a computational perspective, these deficits have been linked to "
        "abnormalities in the hierarchical encoding of precision. Adams and "
        "Friston (2016) proposed that reduced synaptic gain at higher levels of "
        "the cortical hierarchy in schizophrenia leads to imprecise prior "
        "beliefs, which in turn disrupts the balance between predictions and "
        "prediction errors. Valton et al. (2017) further argued that Bayesian "
        "models of schizophrenia remain under-investigated despite their "
        "theoretical promise, calling for more empirical work applying formal "
        "computational models to clinical data."
    )

    add_heading_h2(doc, "1.4 Contribution of This Study")
    add_body_text(doc,
        "The present study addresses this gap by formalizing the corollary "
        "discharge process as a linear Bayesian model and fitting it to EEG "
        "trial data from healthy controls and schizophrenia patients. "
        "Specifically, we: (1) derive a precision weight parameter that "
        "quantifies how strongly the brain integrates motor preparation signals "
        "into auditory predictions, (2) compare this weight between diagnostic "
        "groups to identify compensatory mechanisms, and (3) implement an "
        "ANCOVA-based variance proxy to approximate inter-trial phase coherence "
        "(ITPC) as a measure of efference copy temporal reliability. This work "
        "provides a simple yet theoretically grounded quantitative framework "
        "for characterizing predictive coding deficits from standard EEG "
        "measurements."
    )


def write_methods(doc):
    """Write the methods section."""
    add_heading_h1(doc, "2. Methods")

    add_heading_h2(doc, "2.1 Dataset")
    add_body_text(doc,
        "The dataset used in this study was derived from the EEG experiment "
        "described by Ford et al. (2014). The dataset consists of trial-level "
        "ERP measurements from 81 participants: 32 healthy controls (HC) and "
        "49 patients diagnosed with schizophrenia or schizoaffective disorder "
        "(SZ). Demographic information including age, gender, and education "
        "level was available for all participants."
    )
    add_body_text(doc,
        "The experimental paradigm included three conditions: (1) Active "
        "Button Press + Tone (participants pressed a button to deliver a tone), "
        "(2) Passive Tone Playback (participants heard the same tones played "
        "back without pressing a button), and (3) Button Press Only (participants "
        "pressed a button without an associated tone). Each participant "
        "completed approximately 100 trials per condition."
    )
    add_body_text(doc,
        "ERP amplitudes were extracted at nine scalp electrode sites: Fz, FCz, "
        "Cz, FC3, FC4, C3, C4, CP3, and CP4. For each trial, the following "
        "features were pre-computed: N100 amplitude (auditory response at 100ms "
        "post-stimulus), P200 amplitude (positive deflection at 200ms), and "
        "baseline activity in two windows (B0 and B1, representing pre-stimulus "
        "motor preparation). Trials contaminated by artifacts were flagged "
        "for rejection."
    )

    add_heading_h2(doc, "2.2 Bayesian Model Formulation")
    add_body_text(doc,
        "We conceptualize the corollary discharge process as a linear Bayesian "
        "prediction error model. In this framework, the auditory N100 response "
        "during an active button press represents the posterior — the brain's "
        "residual response after subtracting its prediction from the sensory "
        "input. The model is expressed as:"
    )

    # Equation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("N100_Active = N100_Passive − x × Motor_Prep")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True

    add_body_text(doc, "Where:")
    add_body_text(doc,
        "• N100_Active is the average frontal-central N100 amplitude during "
        "Condition 1 (Active Button Press + Tone), calculated as the mean of "
        "Fz_N100, FCz_N100, and Cz_N100. This represents the Prediction Error "
        "(Posterior) — the residual auditory response after the brain's "
        "prediction has been applied."
    )
    add_body_text(doc,
        "• N100_Passive is the average frontal-central N100 amplitude during "
        "Condition 2 (Passive Tone Playback). This represents the Sensory "
        "Baseline (Likelihood) — the full, unsuppressed auditory response in "
        "the absence of motor prediction."
    )
    add_body_text(doc,
        "• Motor_Prep is the C3_B1 value — the pre-stimulus baseline activity "
        "at electrode C3 (positioned over left motor cortex, contralateral to "
        "the responding right hand) in the B1 time window. This represents the "
        "Efference Copy (Prior) — the motor preparation signal that generates "
        "the sensory prediction."
    )
    add_body_text(doc,
        "• x is the Precision Weight, representing how confidently the brain "
        "integrates (scales) the efference copy signal into an auditory "
        "prediction. Solving for x yields the quantitative strength of the "
        "brain's predictive mechanism."
    )

    add_heading_h2(doc, "2.3 Subject-Level Aggregation")
    add_body_text(doc,
        "To ensure that each participant contributed equally to group-level "
        "statistics regardless of their trial count, we employed subject-level "
        "aggregation (grand mean method). For each participant, the mean values "
        "of N100 (across three frontal-central electrodes) and C3_B1 were first "
        "computed across all non-rejected trials within each condition. "
        "Group-level statistics (HC and SZ) were then calculated as the mean of "
        "these subject-level means. This approach prevents participants with "
        "more trials from disproportionately influencing the group averages."
    )

    add_heading_h2(doc, "2.4 ANCOVA + Variance-Based Proxy for Efference Copy Precision")
    add_body_text(doc,
        "To approximate Inter-Trial Phase Coherence (ITPC) — a measure of the "
        "temporal consistency of the efference copy signal across repeated "
        "trials — we implemented an ANCOVA-based variance proxy. True ITPC "
        "calculation requires raw trial-level time-series data and the Hilbert "
        "transform; however, the available dataset contains only pre-computed "
        "ERP amplitudes per trial. Our proxy method proceeds as follows:"
    )
    add_body_text(doc,
        "Step 1: For each participant, calculate the Mean and Standard Deviation "
        "(SD) of C3_B1 across all Condition 1 trials."
    )
    add_body_text(doc,
        "Step 2: Regress SD on Mean across all participants (pooled across "
        "groups) using ordinary least squares. This step removes the mathematical "
        "coupling between SD and Mean, as higher-amplitude signals naturally "
        "exhibit larger absolute fluctuations."
    )
    add_body_text(doc,
        "Step 3: Extract the residuals from this regression. These residuals — "
        "termed the Adjusted SD — represent the component of trial-to-trial "
        "variability that is not predicted by mean amplitude, serving as a proxy "
        "for the temporal precision (consistency) of the efference copy signal."
    )
    add_body_text(doc,
        "Step 4: Compare the Adjusted SD between HC and SZ groups using an "
        "independent samples t-test, with Cohen's d as the effect size measure."
    )


def write_results(doc):
    """Write the results section."""
    add_heading_h1(doc, "3. Results")

    add_heading_h2(doc, "3.1 Bayesian Precision Weight")
    add_body_text(doc,
        "The precision weight (x) was calculated for both groups using "
        "subject-level aggregation. Table 1 summarizes the key variables and "
        "computed weights."
    )

    add_table(doc,
        headers=["Metric", "HC (N=32)", "SZ (N=49)"],
        rows=[
            ["N100 Active (Cond. 1)", "-3.6384 µV", "-2.6045 µV"],
            ["N100 Passive (Cond. 2)", "-5.6565 µV", "-4.0863 µV"],
            ["Motor Prep (C3_B1)", "1.4425 µV", "0.7856 µV"],
            ["Total Suppression", "2.0181 µV", "1.4818 µV"],
            ["Precision Weight (x)", "-1.3990", "-1.8862"],
        ],
        caption="Table 1. Bayesian Model Parameters by Group"
    )

    add_body_text(doc,
        "The efference copy signal (C3_B1) was approximately 45% weaker in SZ "
        "patients compared to HC (0.79 vs 1.44 µV). This weaker motor "
        "preparation signal translated into reduced total N1 suppression in the "
        "SZ group (1.48 vs 2.02 µV). The resulting linear models are:"
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("HC:  y = c − 1.3990x")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("SZ:  y = c − 1.8862x")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True

    add_body_text(doc,
        "where y = N100_Active, c = N100_Passive, and x = Motor_Prep (C3_B1)."
    )

    add_heading_h2(doc, "3.2 Compensatory Over-Weighting in Schizophrenia")
    add_body_text(doc,
        "A notable finding is that the absolute value of the precision weight "
        "was larger in the SZ group (|−1.89| vs |−1.40|). This indicates that "
        "the schizophrenia brain applies a higher gain to whatever motor "
        "preparation signal it does produce, as though attempting to compensate "
        "for the weakness of the underlying signal. However, because the base "
        "signal (0.79 µV) is substantially degraded, this compensatory "
        "amplification still fails to achieve the same level of N1 suppression "
        "seen in healthy controls."
    )

    add_heading_h2(doc, "3.3 ANCOVA Regression Results")
    add_body_text(doc,
        "The ANCOVA regression confirmed significant mathematical coupling "
        "between SD and Mean of C3_B1 across subjects. Table 2 presents the "
        "regression parameters."
    )

    add_table(doc,
        headers=["Parameter", "Value"],
        rows=[
            ["Regression equation", "SD = 10.9922 + (-0.7913) × Mean"],
            ["R-squared", "0.1398"],
            ["p-value", "0.000586"],
        ],
        caption="Table 2. ANCOVA Regression: SD ~ Mean (Pooled)"
    )

    add_body_text(doc,
        "The regression was highly significant (p < 0.001), confirming that "
        "approximately 14% of the variance in trial-to-trial SD of C3_B1 is "
        "explained by mean amplitude alone, validating the need for ANCOVA "
        "correction."
    )

    add_heading_h2(doc, "3.4 Adjusted SD (ITPC Proxy) Group Comparison")

    add_table(doc,
        headers=["Metric", "HC (N=32)", "SZ (N=49)"],
        rows=[
            ["Mean C3_B1", "1.4425 µV", "0.7856 µV"],
            ["Raw SD", "9.4648 µV", "10.6226 µV"],
            ["Adjusted SD (ITPC Proxy)", "-0.3859 µV", "+0.2520 µV"],
        ],
        caption="Table 3. Efference Copy Precision Metrics by Group"
    )

    add_table(doc,
        headers=["Statistical Test", "Value"],
        rows=[
            ["t-statistic", "-0.7445"],
            ["p-value", "0.4588"],
            ["Significant (p < 0.05)", "No"],
            ["Cohen's d", "-0.1712"],
        ],
        caption="Table 4. Independent t-test: Adjusted SD (HC vs SZ)"
    )

    add_body_text(doc,
        "The direction of the Adjusted SD difference was consistent with "
        "expectations: HC showed negative Adjusted SD (−0.39 µV, indicating "
        "less variability than predicted by amplitude), while SZ showed positive "
        "Adjusted SD (+0.25 µV, indicating more variability). However, this "
        "difference did not reach statistical significance (p = 0.459, Cohen's "
        "d = −0.17), reflecting a small effect size."
    )


def write_discussion(doc):
    """Write the discussion section."""
    add_heading_h1(doc, "4. Discussion")

    add_heading_h2(doc, "4.1 The Efference Copy as a Bayesian Prior")
    add_body_text(doc,
        "Our results demonstrate that the corollary discharge process can be "
        "effectively modeled as a simple linear Bayesian prediction error "
        "system. The model captures the fundamental relationship: the brain's "
        "auditory response to a self-generated sound equals its response to an "
        "externally generated sound minus a prediction term derived from motor "
        "cortex activity. This formalization aligns with the broader predictive "
        "coding framework articulated by Adams and Friston (2016), in which "
        "perception is understood as the minimization of prediction error "
        "weighted by precision."
    )
    add_body_text(doc,
        "The critical finding that the efference copy signal (C3_B1) is "
        "approximately 45% weaker in schizophrenia patients provides "
        "quantitative support for the hypothesis that the predictive mechanism "
        "is structurally impaired at its source — the motor planning phase "
        "(Ford et al., 2014). This is consistent with Pynn and DeSouza's "
        "(2012) review, which emphasized that efference copy signals are "
        "fundamental to distinguishing self-generated from externally generated "
        "sensations, and that failure of this mechanism may underlie auditory "
        "hallucinations."
    )

    add_heading_h2(doc, "4.2 Compensatory Over-Weighting: A Novel Finding")
    add_body_text(doc,
        "Perhaps the most intriguing result is the paradoxically higher "
        "precision weight in the SZ group (|−1.89| vs |−1.40|). Within the "
        "Bayesian framework, this can be interpreted as the schizophrenia brain "
        "attempting to compensate for its degraded efference copy by assigning "
        "disproportionately high precision to whatever motor signal it does "
        "produce. This is analogous to a listener in a noisy room who turns up "
        "the volume — the amplification helps, but if the source signal is too "
        "corrupted, the message remains unintelligible."
    )
    add_body_text(doc,
        "This finding resonates with Fletcher and Frith's (2008) Bayesian "
        "account of positive symptoms, in which they proposed that "
        "schizophrenia involves an aberrant balance between prior expectations "
        "and sensory evidence. Our data suggest that the imbalance may not be "
        "a simple reduction in prior precision, but rather a complex "
        "dysregulation where the brain over-compensates for weak priors, "
        "potentially contributing to the aberrant salience that characterizes "
        "psychotic experiences."
    )

    add_heading_h2(doc, "4.3 Limitations of the ITPC Proxy")
    add_body_text(doc,
        "The non-significant ANCOVA results (p = 0.459) for the Adjusted SD "
        "comparison highlight an important methodological limitation. The "
        "variance-based proxy captures amplitude consistency across trials but "
        "cannot directly measure temporal phase coherence, which is the true "
        "quantity of interest for assessing efference copy timing precision. "
        "True ITPC calculation via the Hilbert transform requires raw "
        "trial-level time-series waveforms, which were not available in the "
        "present dataset."
    )
    add_body_text(doc,
        "Additionally, the sample size (N = 81) may be insufficient to detect "
        "what appears to be a small effect (Cohen's d = −0.17). Future work "
        "with larger samples and access to raw EEG time-series data would "
        "enable direct ITPC calculation and potentially reveal significant "
        "group differences in temporal precision."
    )

    add_heading_h2(doc, "4.4 Additional Limitations")
    add_body_text(doc,
        "Several additional limitations should be noted. First, all patients "
        "in the original dataset were medicated, leaving open the possibility "
        "that observed effects reflect pharmacological rather than disease-"
        "related processes. However, Ford et al. (2014) argued against this "
        "interpretation, citing evidence of corollary discharge deficits in "
        "unmedicated clinical high-risk individuals and first-degree relatives. "
        "Second, the linear model assumes a simple multiplicative relationship "
        "between efference copy strength and prediction weight, which may not "
        "capture nonlinear dynamics present in biological systems. Third, the "
        "use of a single electrode (C3) to index motor preparation may not fully "
        "capture the distributed nature of motor planning activity, as noted "
        "by Valton et al. (2017) in their review of computational approaches "
        "to schizophrenia."
    )


def write_conclusion(doc):
    """Write the conclusion section."""
    add_heading_h1(doc, "5. Conclusion")

    add_body_text(doc,
        "This study formalized the corollary discharge process as a linear "
        "Bayesian prediction error model and applied it to EEG data from "
        "healthy controls and schizophrenia patients. The model successfully "
        "captures the core deficit in schizophrenia: a substantially weakened "
        "efference copy signal that fails to generate adequate sensory "
        "predictions, resulting in reduced N1 suppression to self-generated "
        "sounds."
    )
    add_body_text(doc,
        "The novel finding of compensatory over-weighting in the schizophrenia "
        "group — where the brain applies a disproportionately high precision "
        "weight to its degraded motor signals — offers a new computational "
        "perspective on the mechanisms underlying predictive coding failure in "
        "psychosis. Rather than a simple deficit model, our results suggest "
        "an active but ultimately unsuccessful compensatory process."
    )
    add_body_text(doc,
        "Future work should extend this framework by: (1) incorporating true "
        "ITPC measures using raw EEG time-series data, (2) testing the model "
        "in longitudinal designs to track changes in precision weighting across "
        "illness stages, and (3) examining whether antipsychotic medication "
        "modulates the precision weight parameter. The simplicity and "
        "interpretability of the proposed model make it a promising tool for "
        "translational research bridging computational psychiatry and clinical "
        "neuroscience."
    )


def write_references(doc):
    """Write the references section."""
    add_heading_h1(doc, "References")

    refs = [
        "Adams, R. A., & Friston, K. J. (2016). Brain Computations in "
        "Schizophrenia. In The Neurobiology of Schizophrenia (pp. 283–295). "
        "Elsevier. https://doi.org/10.1016/B978-0-12-801829-3.00024-0",

        "Fletcher, P. C., & Frith, C. D. (2008). Perceiving is believing: a "
        "Bayesian approach to explaining the positive symptoms of schizophrenia. "
        "Nature Reviews Neuroscience, 10, 48–58. "
        "https://doi.org/10.1038/nrn2536",

        "Ford, J. M., Palzes, V. A., Roach, B. J., & Mathalon, D. H. (2014). "
        "Did I Do That? Abnormal Predictive Processes in Schizophrenia When "
        "Button Pressing to Deliver a Tone. Schizophrenia Bulletin, 40(4), "
        "804–812. https://doi.org/10.1093/schbul/sbt072",

        "Pynn, L. K., & DeSouza, J. F. X. (2012). The function of efference "
        "copy signals: Implications for symptoms of schizophrenia. Vision "
        "Research, 76, 124–133. https://doi.org/10.1016/j.visres.2012.10.019",

        "Valton, V., Romaniuk, L., Steele, D., Lawrie, S., & Seriès, P. (2017). "
        "Comprehensive review: Computational modelling of Schizophrenia. "
        "Neuroscience & Biobehavioral Reviews, 83, 631–646. "
        "https://doi.org/10.1016/j.neubiorev.2017.08.022",

        "Friston, K. J. (2005). A theory of cortical responses. Philosophical "
        "Transactions of the Royal Society B, 360(1456), 815–836. "
        "https://doi.org/10.1098/rstb.2005.1622",

        "Crapse, T. B., & Sommer, M. A. (2008). Corollary discharge across "
        "the animal kingdom. Nature Reviews Neuroscience, 9(8), 587–600. "
        "https://doi.org/10.1038/nrn2457",
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"[{i}] {ref}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)


# ============================================================
# MAIN
# ============================================================

def main():
    print("Generating paper_draft.docx...")

    doc = Document()

    # Set default page margins
    for section in doc.sections:
        section.left_margin = Inches(0.88)
        section.right_margin = Inches(0.88)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    # Write all sections
    write_title_block(doc)
    write_abstract(doc)
    write_introduction(doc)
    write_methods(doc)
    write_results(doc)
    write_discussion(doc)
    write_conclusion(doc)
    write_references(doc)

    # Save
    output_path = 'paper_draft.docx'
    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    print(f"Saved: {output_path} ({file_size:,} bytes)")
    print("Done!")


if __name__ == "__main__":
    main()
